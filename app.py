import io
import json
import os
import re
from pathlib import Path

import pandas as pd
import streamlit as st
from PIL import Image, ImageOps, ImageFilter
import pytesseract
import fitz  # PyMuPDF
from pypdf import PdfReader
from docx import Document
from ai_order_parser import ai_parse_order_text, ai_parse_order_image, ai_to_parser_text

st.set_page_config(page_title="Cheese SAP Order Parser", page_icon="🧀", layout="wide")

# Existing master mapping is kept intact.
PRODUCTS = {
"FG-02-0012":{"name":"Classic Cheddar Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["classic cheddar","classic chadder","classic cheddar block"]},
"FG-02-0068":{"name":"Top Cow Cheddar Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["top cow cheddar block","top cow chadder block"]},
"FG-02-0006":{"name":"Achha Pizza Cheddar Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["pizza cheddar","pizza chadder","pizza cheddar block"]},
"FG-02-0018":{"name":"Regular Cheddar Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["regular cheddar block"]},
"FG-02-0028":{"name":"Yellow Slice 1kg","pack":"slice","pcs_ctn":18,"kg":1,"keywords":["yellow slice","orange slice","burger slice","burger/orange","burger cheese","orange cheese"]},
"FG-02-0023":{"name":"White Slice 1kg","pack":"slice","pcs_ctn":18,"kg":1,"keywords":["white slice"]},
"FG-02-0039":{"name":"Jalapeno Cheddar Slice 1kg","pack":"slice","pcs_ctn":18,"kg":1,"keywords":["jalapeno slice","jalapeno cheddar slice"]},
"FG-02-0038":{"name":"Yellow Slice 800gm","pack":"slice800","pcs_ctn":18,"kg":0.8,"keywords":["yellow slice 800","yellow 800"]},
"FG-02-0037":{"name":"White Slice 800gm","pack":"slice800","pcs_ctn":18,"kg":0.8,"keywords":["white slice 800","white 800"]},
"FG-02-0060":{"name":"Top Cow Mozzarella Block White","pack":"block","pcs_ctn":10,"kg":2,"keywords":["top cow mozzarella block","top cow block"]},
"FG-02-0048":{"name":"Top Cow White Dice/Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["top cow white dice","top cow white shred","top cow white shredded","top cow shred white"]},
"FG-02-0049":{"name":"Top Cow Yellow Dice/Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["top cow yellow dice","top cow yellow shred","top cow yellow shredded"]},
"FG-02-0072":{"name":"Classic 70/30 Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["classic 70/30","classic 70.30"]},
"FG-01-0012":{"name":"Classic Mozzarella Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["classic mozzarella block","classic mozz block","classic mozz blk"]},
"FG-02-0036":{"name":"Classic Mozzarella Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["classic mozzarella shred","classic mozzarella shredded","classic shred"]},
"FG-02-0082":{"name":"Classic Mozzarella Shredded DC","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["classic mozzarella shredded dc"]},
"FG-01-0018":{"name":"Danish Mozzarella Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["danish mozzarella block","danish mozz block","danish mozz blk"]},
"FG-01-0030":{"name":"Danish Mozzarella Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["danish mozzarella shred","danish mozzarella shredded"]},
"FG-01-0036":{"name":"Imported/UK Mozzarella Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["imp uk mozzarella","imported mozzarella shred","uk mozzarella shred"]},
"FG-03-0006":{"name":"Imported 70/30 Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["imp 70/30","imported 70/30"]},
"FG-03-0026":{"name":"Imported 70/30 Dice","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["imp 70/30 dice","imported 70/30 dice"]},
"FG-01-0006":{"name":"Achha Mozzarella Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["achha mozzarella block","achha mozz block","achha mozz blk","red mozz blk","red mozzarella block","mozzarella block"]},
"FG-01-0042":{"name":"Achha Mozzarella Shredded White","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["achha shred","achha shredded","achha shared","achha shared white","blue shredd","blue shred","achha mozzarella shred","achha mozzarella shredded"]},
"FG-01-0054":{"name":"Achha Mozzarella Shredded Yellow","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["achha shred yellow","achha shared yellow"]},
"FG-01-0124":{"name":"Achha White Dice","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["achha white dice"]},
"FG-01-0125":{"name":"Achha Yellow Dice","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["achha yellow dice"]},
"FG-03-0018":{"name":"Local 70/30 Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["local 70/30","local 70.30","locl 70/30","lockl 70/30","lockl 70.30"]},
"FG-02-0051":{"name":"New/M3 70/30 Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["new 70/30","new 70.30","m3 70/30","m3 shred","m3 70/30"]},
"FG-03-0024":{"name":"50/50 Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["50/50 shred","50/50 shredded"]},
"FG-01-0066":{"name":"Latina Mozzarella Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["latina mozzarella shred","latina shred"]},
"FG-01-0111":{"name":"Silver Mozzarella Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["silver mozzarella block","silver mozz block","silver"]},
"FG-01-0110":{"name":"Silver Mozzarella Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["silver mozzarella shred"]},
"FG-03-0025":{"name":"Verona Mozzarella Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["verona mozzarella block","verona mozz block"]},
"FG-01-0072":{"name":"Verona Mozzarella Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["verona shred","verona mozzarella shred"]},
"FG-02-0065":{"name":"Pizza Topping Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["pizza topping block"]},
"FG-02-0064":{"name":"Pizza Topping Shredded","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["pizza topping shred"]},
"FG-06-0004":{"name":"Butter White 82 FAT 1kg","pack":"unit","pcs_ctn":1,"kg":1,"keywords":["butter white 82","butter white"]},
"FG-06-0011":{"name":"Butter Yellow 82 FAT 1kg","pack":"unit","pcs_ctn":1,"kg":1,"keywords":["butter yellow 82","butter yellow"]},
"FG-06-0017":{"name":"Butter White 87 Fat 1kg","pack":"unit","pcs_ctn":1,"kg":1,"keywords":["butter white 87"]},
"FG-06-0018":{"name":"Butter Yellow 87 Fat 1kg","pack":"unit","pcs_ctn":1,"kg":1,"keywords":["butter yellow 87"]},
"FG-06-0003":{"name":"Butter White 500gm","pack":"unit","pcs_ctn":1,"kg":0.5,"keywords":["butter white 500"]},
"FG-06-0010":{"name":"Butter Yellow 500gm","pack":"unit","pcs_ctn":1,"kg":0.5,"keywords":["butter yellow 500"]},
"FG-05-0002":{"name":"Desi Ghee Tin 500gm","pack":"unit","pcs_ctn":1,"kg":0.5,"keywords":["desi ghee 500"]},
"FG-05-0011":{"name":"Desi Ghee Tin 1kg","pack":"unit","pcs_ctn":1,"kg":1,"keywords":["desi ghee 1kg","desi ghee"]},
"FG-05-0005":{"name":"Desi Ghee Tin 16kg","pack":"unit","pcs_ctn":1,"kg":16,"keywords":["desi ghee 16kg"]},
"FG-02-0110":{"name":"Nivora Max White Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["max white dice","max w dice"]},
"FG-02-0109":{"name":"Nivora Max White Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["max white shred","max white shredded"]},
"FG-02-0112":{"name":"Nivora Max Yellow Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["max yellow dice"]},
"FG-02-0111":{"name":"Nivora Max Yellow Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["max yellow shred"]},
"FG-02-0102":{"name":"Nivora MF White Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["mf white","mf white dice"]},
"FG-02-0101":{"name":"Nivora MF White Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["mf white shred"]},
"FG-02-0104":{"name":"Nivora MF Yellow Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["mf yellow","mf yellow dice"]},
"FG-02-0103":{"name":"Nivora MF Yellow Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["mf yellow shred"]},
"FG-02-0106":{"name":"Nivora Pro White Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["pro white","pro w","pro white dice"]},
"FG-02-0105":{"name":"Nivora Pro White Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["pro white shred"]},
"FG-02-0108":{"name":"Nivora Pro Yellow Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["pro yellow dice"]},
"FG-02-0107":{"name":"Nivora Pro Yellow Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["pro yellow shred"]},
"FG-02-0118":{"name":"Nivora PT White Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["pt white","pt w","pizza topping white"]},
"FG-02-0117":{"name":"Nivora PT White Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["pt white shred"]},
"FG-02-0120":{"name":"Nivora PT Yellow Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["pt yellow","pt y","pizza topping yellow"]},
"FG-02-0119":{"name":"Nivora PT Yellow Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["pt yellow shred"]},
"FG-02-0114":{"name":"Nivora VF White Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["vf white dice"]},
"FG-02-0113":{"name":"Nivora VF White Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["vf white shred"]},
"FG-02-0116":{"name":"Nivora VF Yellow Dice","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["vf yellow dice"]},
"FG-02-0115":{"name":"Nivora VF Yellow Shredded","pack":"nivora","pcs_ctn":4,"kg":2.5,"keywords":["vf yellow shred"]},
"FG-02-0094":{"name":"Allana Cheddar Cheese Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["allana cheddar block"]},
"FG-02-0096":{"name":"Allana Mozzarella Cheese Block","pack":"block","pcs_ctn":10,"kg":2,"keywords":["allana mozzarella block"]},
"FG-02-0097":{"name":"Allana Mozzarella Cheese Shredded White","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["allana mozzarella shred white","allana mozzarella shred","allana shred"]},
"FG-02-0100":{"name":"Allana Pizza Cheese 70/30 Shredded White","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["allana 70/30","allana pizza 70/30","allana 70/30 shred"]},
"FG-02-0162":{"name":"Allana Mozzarella Block W.Poly","pack":"block","pcs_ctn":10,"kg":2,"keywords":["allana wpoly mozzarella block"]},
"FG-02-0164":{"name":"Allana Mozzarella Shredded White W.Poly","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["allana wpoly mozzarella shred white"]},
"FG-02-0166":{"name":"Allana Mozzarella Shredded Yellow W.Poly","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["allana wpoly mozzarella shred yellow"]},
"FG-02-0174":{"name":"Allana Pizza Cheese 70/30 White W.Poly","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["allana wpoly 70/30"]},
"FG-02-0182":{"name":"Allana Pizza Cheese 50/50 White W.Poly","pack":"regular","pcs_ctn":5,"kg":2,"keywords":["allana wpoly 50/50"]},
}

# ============================================================
# AUGUST-18 ITEM MASTER EXTENSIONS
# These are explicit W.Poly/WP and newly locked mappings.
# WP items are used ONLY when the order explicitly contains WP/W.P/W.Poly.
# ============================================================
PRODUCTS.update({
    "FG-02-0040": {"name":"Silver Cheddar Block", "pack":"block", "pcs_ctn":10, "kg":2, "keywords":["silver cheddar block","silver chadder block"]},
    "FG-02-0067": {"name":"Silver Cheddar Shred", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["silver cheddar shred","silver cheddar shredded","silver chadder shred"]},
    "FG-02-0080": {"name":"New Silver Cheddar Block", "pack":"block", "pcs_ctn":10, "kg":2, "keywords":["new silver cheddar block","new silver chadder block","new silver chadder blk"]},

    # Achha W.Poly
    "FG-01-0123": {"name":"Achha Mozzarella Block W.Poly 2kg", "pack":"block", "pcs_ctn":10, "kg":2, "keywords":["achha mozz block wp","achha mozzarella block wp"]},
    "FG-01-0119": {"name":"Achha Mozzarella Shredded White W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["achha mozz shred wp","achha mozzarella shredded wp"]},
    "FG-01-0120": {"name":"Latina Mozzarella Shredded WP 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["latina shred wp","latina mozzarella shred wp"]},
    "FG-01-0121": {"name":"Local 70/30 Mozzarella/Cheddar Shredded WP 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["local 70/30 wp","local 70/30 mozzarella cheddar wp"]},
    "FG-01-0122": {"name":"Verona Mozzarella Shredded WP 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["verona shred wp","verona mozzarella shred wp"]},
    "FG-01-0126": {"name":"Imported/UK Mozzarella Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["uk shred wp","imported uk shred wp"]},

    # Top Cow W.Poly — Shred and Dice are the same SKU
    "FG-02-0076": {"name":"Top Cow White Mozzarella Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["top cow white dice wp","top cow white shred wp","top cow white shredded wp"]},
    "FG-02-0077": {"name":"Top Cow Yellow Mozzarella Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["top cow yellow dice wp","top cow yellow shred wp","top cow yellow shredded wp"]},
    "FG-02-0079": {"name":"Top Cow Premium White Mozzarella Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["top cow premium white dice wp","top cow premium white shred wp"]},

    # Nivora W.Poly 2kg variants
    "FG-02-0126": {"name":"Nivora Max White Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0125": {"name":"Nivora Max White Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0128": {"name":"Nivora Max Yellow Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0127": {"name":"Nivora Max Yellow Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0093": {"name":"Nivora MF White Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0092": {"name":"Nivora MF White Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0091": {"name":"Nivora MF Yellow Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0090": {"name":"Nivora MF Yellow Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0122": {"name":"Nivora Pro White Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0121": {"name":"Nivora Pro White Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0124": {"name":"Nivora Pro Yellow Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0123": {"name":"Nivora Pro Yellow Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0134": {"name":"Nivora PT White Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0133": {"name":"Nivora PT White Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0136": {"name":"Nivora PT Yellow Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0135": {"name":"Nivora PT Yellow Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0130": {"name":"Nivora VF White Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0129": {"name":"Nivora VF White Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0132": {"name":"Nivora VF Yellow Dice W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},
    "FG-02-0131": {"name":"Nivora VF Yellow Shredded W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":[]},

    # Allana W.Poly
    "FG-02-0160": {"name":"Allana Cheddar Cheese Block W.Poly 2kg", "pack":"block", "pcs_ctn":10, "kg":2, "keywords":["allana cheddar block wp"]},
    "FG-02-0214": {"name":"Allana Cheddar Cheese Slice White W.Poly 800gm", "pack":"slice800", "pcs_ctn":18, "kg":0.8, "keywords":["allana white slice wp 800"]},
    "FG-02-0215": {"name":"Allana Cheddar Cheese Slice Yellow W.Poly 800gm", "pack":"slice800", "pcs_ctn":18, "kg":0.8, "keywords":["allana yellow slice wp 800"]},
    "FG-02-0159": {"name":"Allana Gold Cheddar Cheese Block W.Poly 2kg", "pack":"block", "pcs_ctn":10, "kg":2, "keywords":["allana gold cheddar block wp"]},
    "FG-02-0161": {"name":"Allana Gold Mozzarella Cheese Block W.Poly 2kg", "pack":"block", "pcs_ctn":10, "kg":2, "keywords":["allana gold mozzarella block wp"]},
    "FG-02-0167": {"name":"Allana Gold Mozzarella Cheese Dice White W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold mozz white dice wp"]},
    "FG-02-0169": {"name":"Allana Gold Mozzarella Cheese Dice Yellow W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold mozz yellow dice wp"]},
    "FG-02-0163": {"name":"Allana Gold Mozzarella Cheese Shredded White W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold mozz white shred wp"]},
    "FG-02-0165": {"name":"Allana Gold Mozzarella Cheese Shredded Yellow W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold mozz yellow shred wp"]},
    "FG-02-0177": {"name":"Allana Gold Pizza Cheese Dice 70/30 White W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold pizza dice 70/30 white wp"]},
    "FG-02-0179": {"name":"Allana Gold Pizza Cheese Dice 70/30 Yellow W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold pizza dice 70/30 yellow wp"]},
    "FG-02-0181": {"name":"Allana Gold Pizza Cheese Shredded 50/50 White W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold pizza shred 50/50 white wp"]},
    "FG-02-0183": {"name":"Allana Gold Pizza Cheese Shredded 50/50 Yellow W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold pizza shred 50/50 yellow wp"]},
    "FG-02-0173": {"name":"Allana Gold Pizza Cheese Shredded 70/30 White W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold pizza shred 70/30 white wp"]},
    "FG-02-0175": {"name":"Allana Gold Pizza Cheese Shredded 70/30 Yellow W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana gold pizza shred 70/30 yellow wp"]},
    "FG-02-0171": {"name":"Allana Gold Pizza Topping Block W.Poly 2kg", "pack":"block", "pcs_ctn":10, "kg":2, "keywords":["allana gold pizza topping block wp"]},
    "FG-02-0168": {"name":"Allana Mozzarella Cheese Dice White W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana mozz white dice wp"]},
    "FG-02-0170": {"name":"Allana Mozzarella Cheese Dice Yellow W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana mozz yellow dice wp"]},
    "FG-02-0178": {"name":"Allana Pizza Cheese Dice 70/30 White W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana pizza dice 70/30 white wp"]},
    "FG-02-0180": {"name":"Allana Pizza Cheese Dice 70/30 Yellow W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana pizza dice 70/30 yellow wp"]},
    "FG-02-0184": {"name":"Allana Pizza Cheese Shredded 50/50 Yellow W.Poly 2kg", "pack":"regular", "pcs_ctn":5, "kg":2, "keywords":["allana pizza shred 50/50 yellow wp"]},
    "FG-02-0172": {"name":"Allana Pizza Topping Block W.Poly 2kg", "pack":"block", "pcs_ctn":10, "kg":2, "keywords":["allana pizza topping block wp"]},
})

RULES_FILE = Path("rules.json")

def norm(s):
    s=str(s).lower().strip()
    replacements={
        "shared":"shred","shraded":"shredded","shrad":"shred","shrd":"shred","shreded":"shredded",
        "chadder":"cheddar","chaddar":"cheddar","cheder":"cheddar","cheedar":"cheddar","chesse":"cheese",
        "accha":"achha","acha":"achha","locl":"local","lockl":"local","70.30":"70/30","70:30":"70/30",
        "mozzrela":"mozzarella","mozzarela":"mozzarella","mozralla":"mozzarella","blk":"block",
        "coton":"ctn","cotton":"ctn","cartons":"ctn","carton":"ctn",
        "packets":"pkt","packet":"pkt","pac":"pkt","pcs":"pkt"
    }
    for a,b in replacements.items(): s=s.replace(a,b)
    return re.sub(r"\s+"," ",s)

def load_rules():
    default={"product_aliases":[],"quantity_rules":[],"customer_rules":[],"general_rules":[]}
    try:
        if RULES_FILE.exists():
            with RULES_FILE.open("r",encoding="utf-8") as f: return {**default,**json.load(f)}
    except Exception: pass
    return default

def save_rules(rules):
    with RULES_FILE.open("w",encoding="utf-8") as f: json.dump(rules,f,ensure_ascii=False,indent=2)
    return True

RULES=load_rules()

def apply_saved_aliases(text):
    t=norm(text)
    for rule in RULES.get("product_aliases",[]):
        alias=norm(rule.get("alias",""));code=rule.get("code","")
        if alias and code in PRODUCTS and alias in t:return code
    return None

def find_product(text):
    """
    Deterministic product mapper.
    IMPORTANT:
    - WP/W.P/W.Poly codes are used ONLY when WP is explicitly written.
    - Without explicit WP, regular/non-WP mapping is used.
    - Specific identifiers (Danish, New Silver, Silver Cheddar, Pizza Topping, etc.)
      always override generic defaults.
    """
    t=norm(text)
    saved=apply_saved_aliases(t)
    if saved:return saved

    # Explicit WP/W.Poly detector. Do NOT infer WP.
    wp = bool(re.search(r"(?:\bwp\b|w\.\s*p\.?\b|\bw[\s-]?poly\b|\bw/p\b)", t, re.I))

    # ---------- BUTTER: SALTED = YELLOW, UNSALTED = WHITE ----------
    if "butter" in t:
        salted = bool(re.search(r"\bsalted\b", t))
        unsalted = bool(re.search(r"\bunsalted\b", t))
        has_yellow = "yellow" in t
        has_white = "white" in t
        if salted and has_white and not has_yellow:
            return None  # explicit conflict => review
        if unsalted and has_yellow and not has_white:
            return None  # explicit conflict => review
        yellow = salted or has_yellow
        white = unsalted or has_white
        fat87 = "87" in t
        gm500 = "500" in t or "0.5" in t
        if yellow:
            if fat87:return "FG-06-0018"
            if gm500:return "FG-06-0010"
            return "FG-06-0011"
        if white:
            if fat87:return "FG-06-0017"
            if gm500:return "FG-06-0003"
            return "FG-06-0004"
        return None  # no safe salted/unsalted/color signal

    # ---------- EXPLICIT WP MAPPINGS ----------
    if wp:
        # Top Cow: Shred and Dice are the SAME SKU.
        if "top cow" in t:
            if "premium" in t and "white" in t:
                return "FG-02-0079"
            if "yellow" in t:
                return "FG-02-0077"
            if "white" in t or "dice" in t or "shred" in t:
                return "FG-02-0076"

        if "new silver" in t and "cheddar" in t and "block" in t:
            return "FG-02-0080"
        if "silver" in t and "cheddar" in t and "shred" in t:
            return "FG-02-0067"
        if "silver" in t and "cheddar" in t and "block" in t:
            return "FG-02-0040"

        if "achha" in t and "block" in t and "mozz" in t:
            return "FG-01-0123"
        if "achha" in t and "shred" in t and ("mozz" in t or "achha" in t):
            return "FG-01-0119"

        # Allana WP / W.Poly — specific qualifiers first.
        if "allana" in t:
            if "gold" in t and "cheddar" in t and "block" in t:return "FG-02-0159"
            if "gold" in t and "mozz" in t and "block" in t:return "FG-02-0161"
            if "gold" in t and "mozz" in t and "dice" in t and "yellow" in t:return "FG-02-0169"
            if "gold" in t and "mozz" in t and "dice" in t:return "FG-02-0167"
            if "gold" in t and "mozz" in t and "shred" in t and "yellow" in t:return "FG-02-0165"
            if "gold" in t and "mozz" in t and "shred" in t:return "FG-02-0163"
            if "gold" in t and "pizza" in t and "70/30" in t and "dice" in t and "yellow" in t:return "FG-02-0179"
            if "gold" in t and "pizza" in t and "70/30" in t and "dice" in t:return "FG-02-0177"
            if "gold" in t and "pizza" in t and "50/50" in t and "yellow" in t:return "FG-02-0183"
            if "gold" in t and "pizza" in t and "50/50" in t and "white" in t:return "FG-02-0181"
            if "gold" in t and "pizza" in t and "70/30" in t and "yellow" in t:return "FG-02-0175"
            if "gold" in t and "pizza" in t and "70/30" in t:return "FG-02-0173"
            if "gold" in t and "pizza topping" in t and "block" in t:return "FG-02-0171"
            if "cheddar" in t and "slice" in t and "white" in t:return "FG-02-0214"
            if "cheddar" in t and "slice" in t and "yellow" in t:return "FG-02-0215"
            if "cheddar" in t and "block" in t:return "FG-02-0160"
            if "pizza topping" in t and "block" in t:return "FG-02-0172"
            if "pizza" in t and "70/30" in t and "dice" in t and "yellow" in t:return "FG-02-0180"
            if "pizza" in t and "70/30" in t and "dice" in t:return "FG-02-0178"
            if "pizza" in t and "70/30" in t and "shred" in t and "yellow" in t:return "FG-02-0176"
            if "pizza" in t and "70/30" in t and "shred" in t:return "FG-02-0174"
            if "pizza" in t and "50/50" in t and "yellow" in t:return "FG-02-0184"
            if "pizza" in t and "50/50" in t and "white" in t:return "FG-02-0182"
            if "mozz" in t and "dice" in t and "yellow" in t:return "FG-02-0170"
            if "mozz" in t and "dice" in t:return "FG-02-0168"
            if "mozz" in t and "shred" in t and "yellow" in t:return "FG-02-0166"
            if "mozz" in t and "shred" in t:return "FG-02-0164"
            if "mozz" in t and "block" in t:return "FG-02-0162"

        # Nivora WP: weight + family + color + form are mandatory when both weights exist.
        nivora_families = [("mf","009"),("pro","012"),("max","012"),("pt","013"),("vf","013")]
        weight25 = bool(re.search(r"\b2\.5\s*kg\b|\b2\.5kg\b", t))
        weight2 = bool(re.search(r"\b2\s*kg\b|\b2kg\b", t))
        color = "yellow" if "yellow" in t else "white" if "white" in t else None
        form = "dice" if "dice" in t else "shred" if "shred" in t else None
        if any(f" {fam} " in f" {t} " for fam,_ in nivora_families):
            fam = next(fam for fam,_ in nivora_families if re.search(rf"\b{fam}\b", t))
            table={
                ("mf","white","dice",2):"FG-02-0093",("mf","white","shred",2):"FG-02-0092",
                ("mf","yellow","dice",2):"FG-02-0091",("mf","yellow","shred",2):"FG-02-0090",
                ("pro","white","dice",2):"FG-02-0122",("pro","white","shred",2):"FG-02-0121",
                ("pro","yellow","dice",2):"FG-02-0124",("pro","yellow","shred",2):"FG-02-0123",
                ("max","white","dice",2):"FG-02-0126",("max","white","shred",2):"FG-02-0125",
                ("max","yellow","dice",2):"FG-02-0128",("max","yellow","shred",2):"FG-02-0127",
                ("pt","white","dice",2):"FG-02-0134",("pt","white","shred",2):"FG-02-0133",
                ("pt","yellow","dice",2):"FG-02-0136",("pt","yellow","shred",2):"FG-02-0135",
                ("vf","white","dice",2):"FG-02-0130",("vf","white","shred",2):"FG-02-0129",
                ("vf","yellow","dice",2):"FG-02-0132",("vf","yellow","shred",2):"FG-02-0131",
            }
            table25={
                ("mf","white","dice",2.5):"FG-02-0102",("mf","white","shred",2.5):"FG-02-0101",
                ("mf","yellow","dice",2.5):"FG-02-0104",("mf","yellow","shred",2.5):"FG-02-0103",
                ("pro","white","dice",2.5):"FG-02-0106",("pro","white","shred",2.5):"FG-02-0105",
                ("pro","yellow","dice",2.5):"FG-02-0108",("pro","yellow","shred",2.5):"FG-02-0107",
                ("max","white","dice",2.5):"FG-02-0110",("max","white","shred",2.5):"FG-02-0109",
                ("max","yellow","dice",2.5):"FG-02-0112",("max","yellow","shred",2.5):"FG-02-0111",
                ("pt","white","dice",2.5):"FG-02-0118",("pt","white","shred",2.5):"FG-02-0117",
                ("pt","yellow","dice",2.5):"FG-02-0120",("pt","yellow","shred",2.5):"FG-02-0119",
                ("vf","white","dice",2.5):"FG-02-0114",("vf","white","shred",2.5):"FG-02-0113",
                ("vf","yellow","dice",2.5):"FG-02-0116",("vf","yellow","shred",2.5):"FG-02-0115",
            }
            key=(fam,color,form)
            if weight25 and key in table25:return table25[key]
            if weight2 and key in table:return table[key]
            if not weight2 and not weight25:return None  # ambiguous => review

        # Other explicit WP regulars.
        if "latina" in t and "shred" in t:return "FG-01-0120"
        if "local" in t and "70/30" in t and "shred" in t:return "FG-01-0121"
        if "verona" in t and "shred" in t:return "FG-01-0122"
        if ("imported" in t or "uk" in t) and "shred" in t:return "FG-01-0126"

    # ---------- NON-WP / REGULAR MAPPINGS ----------
    if "new silver" in t and "cheddar" in t and "block" in t:return "FG-02-0080"
    if "silver" in t and "cheddar" in t and "shred" in t:return "FG-02-0067"
    if "silver" in t and "cheddar" in t and "block" in t:return "FG-02-0040"

    if "50/50" in t and ("shred" in t or "shredded" in t):return "FG-03-0024"
    if "classic" in t and ("shred" in t or "shredded" in t) and "70/30" not in t:
        if "dc" in t: return "FG-02-0082"
        return "FG-02-0036"

    priority=[
        ("red mozz blk","FG-01-0006"),("red mozzarella block","FG-01-0006"),
        ("blue shredd","FG-01-0042"),("blue shred","FG-01-0042"),
        ("danish mozzarella block","FG-01-0018"),("danish mozz block","FG-01-0018"),
        ("classic mozzarella block","FG-01-0012"),("classic mozz block","FG-01-0012"),
        ("burger slice","FG-02-0028"),("orange slice","FG-02-0028")
    ]
    for key,code in priority:
        if key in t:return code

    # Slice selection: explicit 800gm wins; otherwise 1kg.
    if "slice" in t:
        white="white" in t; yellow=any(x in t for x in ("yellow","orange","burger")); jal="jalapeno" in t or "jal" in t
        is800=any(x in t for x in ("800","800gm","800 gm",".8","0.8"))
        if jal:return "FG-02-0039"
        if white:return "FG-02-0037" if is800 else "FG-02-0023"
        if yellow:return "FG-02-0038" if is800 else "FG-02-0028"

    if "local" in t and "70/30" in t:return "FG-03-0018"
    if ("new" in t or "m3" in t) and "70/30" in t:return "FG-02-0051"
    if "imp" in t and "70/30" in t:return "FG-03-0006"
    if "classic" in t and "70/30" in t:return "FG-02-0072"

    if "achha" in t and "yellow" in t and "dice" in t:return "FG-01-0125"
    if "achha" in t and "white" in t and "dice" in t:return "FG-01-0124"
    if "achha" in t and "yellow" in t and "shred" in t:return "FG-01-0054"
    if "achha" in t and "shred" in t:return "FG-01-0042"

    # Top Cow REGULAR: shred and dice are the SAME SKU.
    if "top cow" in t and "yellow" in t and ("dice" in t or "shred" in t):return "FG-02-0049"
    if "top cow" in t and ("white" in t or "dice" in t or "shred" in t):return "FG-02-0048"
    if "top cow" in t and "cheddar" in t and "block" in t:return "FG-02-0068"
    if "top cow" in t and "block" in t:return "FG-02-0060"

    if "verona" in t and "block" in t:return "FG-03-0025"
    if "verona" in t:return "FG-01-0072"
    if "silver" in t and "shred" in t:return "FG-01-0110"
    if "silver" in t:return "FG-01-0111"
    if "danish" in t and "shred" in t:return "FG-01-0030"
    if "danish" in t:return "FG-01-0018"
    if "classic" in t and "cheddar" in t and "block" in t:return "FG-02-0012"
    if "pizza topping" in t and "block" in t:return "FG-02-0065"
    if "pizza topping" in t and "shred" in t:return "FG-02-0064"
    if "pizza cheddar" in t or ("pizza" in t and "cheddar" in t and "block" in t):return "FG-02-0006"

    for key,code in [
        ("mf white dice","FG-02-0102"),("mf white","FG-02-0102"),
        ("mf yellow dice","FG-02-0104"),("mf yellow","FG-02-0104"),
        ("pro white","FG-02-0106"),("pro w","FG-02-0106"),
        ("pro yellow","FG-02-0108"),
        ("max white","FG-02-0110"),("max yellow","FG-02-0112"),
        ("pt white","FG-02-0118"),("pt w","FG-02-0118"),
        ("pt yellow","FG-02-0120"),("pt y","FG-02-0120"),
        ("vf white","FG-02-0114"),("vf yellow","FG-02-0116")
    ]:
        if key in t:return code

    if "mozzarella block" in t or "mozz block" in t or "mozz blk" in t:return "FG-01-0006"

    candidates = []
    for code, p in PRODUCTS.items():
        is_wp_product = "w.p" in p["name"].lower() or "wp" in p["name"].lower()
        if not wp and is_wp_product:
            continue
        if wp and not is_wp_product:
            continue
        for k in p["keywords"]:
            if k and all(word in t for word in k.split()):
                candidates.append((len(k), code))
    return sorted(candidates, reverse=True)[0][1] if candidates else None

def parse_quantity(line):
    t=norm(line)
    m=re.search(r"(\d+(?:\.\d+)?)\s*kg\s+.*burger\s*(?:/|or)?\s*(?:orange)?\s*slice",t)
    if m:return int(float(m.group(1))),"PKT"
    
    masked = re.sub(r'\b(70/30|50/50|82|87|800\s*gm|800|500\s*gm|500)\b', 'XXX', t)
    
    matches = re.findall(r"(\d+(?:\.\d+)?)\s*(ctn|carton|cartons|box|boxes|pkt|packet|packets|pcs|pc|units?|kg)\b", masked)
    if matches:
        non_weight = [m for m in matches if not (m[1] == 'kg' and m[0] in ('1', '2', '2.5', '1.0', '2.0'))]
        best_match = non_weight[0] if non_weight else matches[0]
        qty = float(best_match[0]); unit = best_match[1].lower(); qty = int(qty) if qty.is_integer() else qty
        if unit in ("ctn","carton","cartons","box","boxes"):return qty,"CTN"
        if unit in ("pkt","packet","packets","pcs","pc","unit","units"):return qty,"PKT"
        if unit=="kg":return qty,"KG"
        
    m = re.search(r"(?<!\d\.)\b(\d+(?:\.\d+)?)\b", masked)
    if m:
        qty=float(m.group(1)); qty=int(qty) if qty.is_integer() else qty
        return qty,"CTN" if qty>10 else "PKT"
        
    return None,None

def is_customer_identifier(line):
    """Return True for standalone customer/BP identifiers; never use these for product mapping."""
    t = norm(line).strip()
    if not t:
        return False
    # Common customer/BP formats: CFS-12345, CFS12345, BP-12345, BP12345.
    return bool(re.fullmatch(r"(?:cfs|bp)[\s_-]*[0-9]{3,}", t, re.I))

def parse_order(text, progress=None):
    lines=[x.strip() for x in str(text).splitlines() if x.strip()]; rows=[]; customer=""
    if progress: progress(0.05,"Reading order text…")
    for i,line in enumerate(lines):
        if i==0 and not re.search(r"\d",line):
            customer=line.strip(); continue
        if is_customer_identifier(line):
            continue
        code=find_product(line); qty,unit=parse_quantity(line)
        if not code or qty is None:
            rows.append({"Source":line,"Customer":customer,"FG Code":"UNMAPPED","Product":"","Input Qty":qty,"Input Unit":unit,"SAP Qty (PKT)":"","Status":"CHECK MAPPING"}); continue
        p=PRODUCTS[code]; sap_qty=round(qty/p["kg"]) if unit=="KG" else int(qty*p["pcs_ctn"]) if unit=="CTN" else int(qty)
        rows.append({"Source":line,"Customer":customer,"FG Code":code,"Product":p["name"],"Input Qty":qty,"Input Unit":unit,"SAP Qty (PKT)":sap_qty,"Status":"OK"})
        if progress: progress(min(0.70,0.10+0.55*(i+1)/max(1,len(lines))),f"Mapping line {i+1}/{len(lines)}…")
    if progress: progress(0.75,"Running quantity and SAP validation…")
    return customer,rows

def sap_line(code,qty):return f"{code}\t\t{int(qty)}\t\t\t\t\tHO-WH\t\tCHEESE"

def extract_excel(file):
    xls=pd.ExcelFile(file); frames=[]
    for sheet in xls.sheet_names:
        df=pd.read_excel(file,sheet_name=sheet)
        if not df.empty:df["__sheet__"]=sheet;frames.append(df)
    return pd.concat(frames,ignore_index=True) if frames else pd.DataFrame()

def normalize_columns(df):
    d=df.copy();d.columns=[str(c).strip().lower() for c in d.columns];return d

def process_tabular(df):
    d=normalize_columns(df);code_col=next((c for c in d.columns if "sap" in c and "code" in c),None);product_col=next((c for c in d.columns if "product" in c or "description" in c or "item" in c),None);ctn_col=next((c for c in d.columns if "carton" in c or c=="ctn"),None);unit_col=next((c for c in d.columns if c in ("units","unit","pcs","qty","quantity")),None);customer_col=next((c for c in d.columns if "customer" in c or "party" in c),None)
    output=[]
    for _,r in d.iterrows():
        customer=str(r.get(customer_col," ")).strip() if customer_col else "";code=str(r.get(code_col," ")).strip() if code_col else "";product=str(r.get(product_col," ")).strip() if product_col else ""
        if code not in PRODUCTS:code=find_product(product)
        if not code:continue
        try:units=float(r.get(unit_col,0)) if unit_col else 0
        except Exception:units=0
        try:ctn=float(r.get(ctn_col,0)) if ctn_col else 0
        except Exception:ctn=0
        if units>0:sap_qty=int(units)
        elif ctn>0:sap_qty=int(ctn*PRODUCTS[code]["pcs_ctn"])
        else:continue
        output.append({"Customer":customer,"FG Code":code,"Product":PRODUCTS[code]["name"],"SAP Qty (PKT)":sap_qty,"Source":"Excel","Status":"OK"})
    return pd.DataFrame(output)

IMAGE_EXTS={"png","jpg","jpeg","webp","bmp","tif","tiff"};TABULAR_EXTS={"xlsx","xls","csv"};TEXT_EXTS={"txt"};PDF_EXTS={"pdf"};DOC_EXTS={"docx"}

def clean_ocr_text(text):
    text=text.replace("\x00"," ");lines=[]
    for raw in text.splitlines():
        line=re.sub(r"[ \t]+"," ",raw).strip()
        if line:lines.append(line)
    return "\n".join(lines)

def ocr_image(data):
    image=Image.open(io.BytesIO(data))
    if image.mode not in ("RGB","L"):image=image.convert("RGB")
    w,h=image.size;scale=2 if max(w,h)<2500 else 1
    if scale>1:image=image.resize((w*scale,h*scale))
    gray=ImageOps.grayscale(image);gray=ImageOps.autocontrast(gray);gray=gray.filter(ImageFilter.SHARPEN)
    text=pytesseract.image_to_string(gray,config="--psm 6")
    if not text.strip():text=pytesseract.image_to_string(gray,config="--psm 11")
    return clean_ocr_text(text)

def pdf_to_text(data):
    parts=[];reader=PdfReader(io.BytesIO(data))
    for page in reader.pages:
        try:parts.append(page.extract_text() or "")
        except Exception:parts.append("")
    text=clean_ocr_text("\n".join(parts))
    if len(re.sub(r"\s","",text))<20:
        doc=fitz.open(stream=data,filetype="pdf");ocr_parts=[]
        for page in doc:
            pix=page.get_pixmap(matrix=fitz.Matrix(2,2),alpha=False);ocr_parts.append(ocr_image(pix.tobytes("png")))
        doc.close();text=clean_ocr_text("\n".join(ocr_parts))
    return text

def docx_to_text(data):
    doc=Document(io.BytesIO(data));chunks=[p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:chunks.append(" ".join(cell.text.strip() for cell in row.cells))
    return clean_ocr_text("\n".join(chunks))

def read_uploaded_file(uploaded_file):
    ext=Path(uploaded_file.name).suffix.lower().lstrip(".");data=uploaded_file.getvalue()
    if ext in IMAGE_EXTS:return "image",ocr_image(data)
    if ext in PDF_EXTS:return "pdf",pdf_to_text(data)
    if ext in DOC_EXTS:return "docx",docx_to_text(data)
    if ext=="csv":return "csv",pd.read_csv(io.BytesIO(data))
    if ext in ("xlsx","xls"):return "excel",extract_excel(io.BytesIO(data))
    if ext in TEXT_EXTS:return "text",data.decode("utf-8",errors="ignore")
    raise ValueError(f"Unsupported file type: .{ext}")

def show_order_result(customer,rows,source_label):
    ok=[r for r in rows if r["Status"]=="OK"];bad=[r for r in rows if r["Status"]!="OK"]
    if customer:st.markdown(f"### {customer}")
    st.subheader("SAP Paste Format");merged={}
    for r in ok:merged[r["FG Code"]]=merged.get(r["FG Code"],0)+int(r["SAP Qty (PKT)"])
    if merged:
        sap="\n".join(sap_line(c,q) for c,q in merged.items());st.code(sap,language="text");st.download_button("Download SAP Order",sap,file_name=f"{customer or 'order'}_SAP.txt",mime="text/plain",key=f"dl_{abs(hash(source_label))}")
    else:st.warning("No mapped order lines found.")
    if bad:st.error("These lines need manual checking:");st.dataframe(pd.DataFrame(bad),use_container_width=True)
    check_df=pd.DataFrame(ok)
    if not check_df.empty:st.subheader("Triple-check");st.dataframe(check_df[["Source","FG Code","Product","Input Qty","Input Unit","SAP Qty (PKT)"]],use_container_width=True)

def teach_rule(rule_text, category="general", customer=""):
    text=rule_text.strip()
    if not text:return False,"Write a teaching rule first."
    rules=load_rules()
    m=re.search(r"(.+?)\s*(?:means|=|is)\s*(FG-\d{2}-\d{4})\b",text,re.I)
    if m and m.group(2).upper() in PRODUCTS:
        alias=m.group(1).strip(" :-");code=m.group(2).upper()
        rules.setdefault("product_aliases",[])
        rules["product_aliases"]=[r for r in rules["product_aliases"] if norm(r.get("alias"))!=norm(alias)]
        rules["product_aliases"].append({"alias":alias,"code":code,"note":text})
        save_rules(rules);return True,f"Saved product alias: {alias} → {code}"
    entry={"rule":text,"customer":customer.strip(),"category":category}
    key="customer_rules" if customer.strip() or category=="customer" else ("quantity_rules" if category=="quantity / conversion" else "general_rules")
    rules.setdefault(key,[]).append(entry);save_rules(rules)
    return True,"Teaching rule saved and will be applied as persistent guidance."

st.title("🧀 Cheese Order Parser → SAP")
st.caption("Smart input + persistent teaching. Press Enter to start processing. Drag/drop WhatsApp screenshots, images, PDFs, Excel, CSV, Word files or paste an order.")

with st.sidebar:
    st.header("🧠 Parser Memory")
    st.metric("Saved product aliases",len(RULES.get("product_aliases",[])))
    st.metric("Saved rules",len(RULES.get("general_rules",[]))+len(RULES.get("customer_rules",[]))+len(RULES.get("quantity_rules",[])))
    st.caption("Saved rules are stored in rules.json. They remain available across normal Streamlit reruns.")
    if RULES.get("product_aliases"):
        st.subheader("Learned aliases")
        for r in RULES["product_aliases"][-12:]:st.write(f"• **{r['alias']}** → `{r['code']}`")
    st.divider()
    st.write("• SAP quantity = PKT/Units")
    st.write("• CTN is converted automatically")
    st.write("• 2kg block = 10 PKT/CTN")
    st.write("• 2kg shred/dice = 5 PKT/CTN")
    st.write("• Nivora 2.5kg = 4 PKT/CTN")
    st.write("• Slice = 18 PKT/CTN")
    st.write("• Red mozz blk = Achha Mozz Block")
    st.write("• Blue shredd = Achha Mozz Shredded")
    st.write("• Danish Mozz Block = Danish, never Achha")
    st.write("• Salted Butter = Yellow; Unsalted Butter = White")
    st.write("• WP code ONLY when WP/W.P/W.Poly is explicitly written")
    st.write("• Top Cow White Shred = Top Cow White Dice (same SKU)")
    st.write("• Top Cow Yellow Shred = Top Cow Yellow Dice (same SKU)")
    st.write("• 50/50 Shredded = Imported 50/50 Shredded 2kg")
    st.write("• Classic Shredded = FG-02-0036")

order_tab,teach_tab,excel_tab=st.tabs(["🤖 Smart Order Input","🎓 Teach / Save Rule","📊 Excel / CSV"])

with order_tab:
    st.subheader("Smart Order Input")
    st.caption("One box for text, WhatsApp screenshots, images, Excel, PDF, CSV, Word and multiple files.")

    submission = st.chat_input(
        "Type/paste WhatsApp order — Ctrl+V screenshot or drag & drop files",
        accept_file="multiple",
        file_type=sorted(IMAGE_EXTS | TABULAR_EXTS | TEXT_EXTS | PDF_EXTS | DOC_EXTS),
        max_upload_size=200,
        key="smart_order_input",
    )

    if submission is not None:
        message = (getattr(submission, "text", "") or "").strip()
        files = list(getattr(submission, "files", []) or [])
        progress = st.progress(0, text="Starting parser…")
        status = st.empty()
        if message:
            try:
                status.info("🧠 AI is understanding the order…")
                ai_result = ai_parse_order_text(message)
                with st.expander("🔎 AI interpretation", expanded=False): st.json(ai_result)
                orders = ai_result.get("orders") or []
                if orders:
                    for order_no, order in enumerate(orders, 1):
                        customer_name = str(order.get("customer_name", "")).strip()
                        item_lines = [customer_name] if customer_name else []
                        for item in order.get("items", []):
                            q = item.get("quantity"); unit = item.get("unit", "PKT"); product = item.get("product", "")
                            if product and q not in (None, ""):
                                qf = float(q)
                                qtxt = str(int(qf)) if qf.is_integer() else str(qf)
                                item_lines.append(f"{qtxt} {unit} {product}")
                        if len(item_lines) > 1:
                            progress.progress(min(0.45 + 0.35*order_no/max(1,len(orders)), 0.90), text=f"Mapping customer {order_no}/{len(orders)}…")
                            customer, rows = parse_order("\n".join(item_lines), lambda v,t: progress.progress(min(0.45 + v*0.35, 0.90), text=t))
                            show_order_result(customer_name or customer, rows, f"chat_text_{order_no}")
                else:
                    parser_text = ai_to_parser_text(ai_result)
                    if parser_text.strip():
                        customer, rows = parse_order(parser_text, lambda v,t: progress.progress(min(0.45 + v*0.45, 0.90), text=t))
                        show_order_result(ai_result.get("customer_name") or customer, rows, "chat_text")
                    else:
                        st.warning("No readable order was detected.")
            except Exception as e: st.error(f"Text parsing error: {e}")
        for idx, f in enumerate(files, 1):
            st.markdown("---")
            st.markdown(f"### 📎 {f.name}")
            try:
                status.info(f"Reading {f.name}…")
                kind, content = read_uploaded_file(f)
                if kind in ("excel", "csv"):
                    result = process_tabular(content)
                    if result.empty: st.warning(f"{f.name}: no valid mapped rows found.")
                    else:
                        st.success(f"{f.name}: {len(result)} mapped rows extracted.")
                        st.dataframe(result, use_container_width=True)
                        for customer_name, group in result.groupby("Customer", dropna=False):
                            cust = str(customer_name).strip()
                            st.markdown(f"### {cust or 'Order'}")
                            merged = {}
                            for _, r in group.iterrows(): merged[r["FG Code"]] = merged.get(r["FG Code"], 0) + int(r["SAP Qty (PKT)"])
                            sap = "\n".join(sap_line(c,q) for c,q in merged.items())
                            st.code(sap, language="text")
                            st.download_button("📋 Copy/Download SAP Order", sap, file_name=f"{cust or 'order'}_SAP.txt", mime="text/plain", key=f"sap_{idx}_{abs(hash(cust))}")
                else:
                    ai_result = ai_parse_order_image(f.getvalue()) if kind == "image" else ai_parse_order_text(str(content))
                    with st.expander("🔎 AI interpretation", expanded=False): st.json(ai_result)
                    orders = ai_result.get("orders") or []
                    if orders:
                        for order_no, order in enumerate(orders, 1):
                            customer_name = str(order.get("customer_name", "")).strip()
                            item_lines = [customer_name] if customer_name else []
                            for item in order.get("items", []):
                                q = item.get("quantity"); unit = item.get("unit", "PKT"); product = item.get("product", "")
                                if product and q not in (None, ""):
                                    qf = float(q)
                                    qtxt = str(int(qf)) if qf.is_integer() else str(qf)
                                    item_lines.append(f"{qtxt} {unit} {product}")
                            if len(item_lines) > 1:
                                progress.progress(min(0.50 + 0.35*order_no/max(1,len(orders)), 0.92), text=f"Mapping customer {order_no}/{len(orders)}…")
                                customer, rows = parse_order("\n".join(item_lines), lambda v,t: progress.progress(min(0.50 + v*0.35, 0.92), text=t))
                                show_order_result(customer_name or customer, rows, f"{f.name}_{order_no}")
                    else:
                        extracted = ai_to_parser_text(ai_result)
                        if extracted.strip():
                            customer, rows = parse_order(extracted, lambda v,t: progress.progress(min(0.50 + v*0.40, 0.92), text=t))
                            show_order_result(ai_result.get("customer_name") or customer, rows, f.name)
                        else:
                            st.warning(f"{f.name}: no readable order was detected.")
            except Exception as e: st.error(f"{f.name}: {e}")
        progress.progress(1.0, text="Done")

with teach_tab:
    st.subheader("🎓 Teach the Parser")
    st.write("Write a rule once and save it. Product aliases can be written naturally, for example: **'mf white dice = FG-02-0102'**.")
    teach_text=st.text_area("Teaching / rule",height=180,placeholder="Example: Customer ABC calls Nivora MF White Dice 'MF White'. Use FG-02-0102.",key="teach_text")
    c1,c2=st.columns([1,1])
    with c1:teach_category=st.selectbox("Rule type",["general","customer","product alias","quantity / conversion"])
    with c2:teach_customer=st.text_input("Customer (optional)",placeholder="Only if this rule is customer-specific")
    if st.button("💾 Save & Remember Rule",type="primary"):
        category="customer" if teach_category=="customer" else ("general" if teach_category=="product alias" else teach_category)
        ok,msg=teach_rule(teach_text,category,teach_customer)
        if ok:st.success(msg);RULES=load_rules()
        else:st.warning(msg)
    st.divider();st.subheader("Saved memory")
    rules=load_rules();aliases=rules.get("product_aliases",[])
    if aliases:st.dataframe(pd.DataFrame(aliases),use_container_width=True)
    for label,key in [("General rules","general_rules"),("Customer rules","customer_rules"),("Quantity / conversion rules","quantity_rules")]:
        items=rules.get(key,[])
        if items:
            st.markdown(f"**{label}**")
            for i,r in enumerate(items,1):st.write(f"{i}. {r.get('rule','')}" + (f" — Customer: {r.get('customer')}" if r.get('customer') else ""))

with excel_tab:
    st.subheader("Upload Excel / CSV")
    uploaded=st.file_uploader("Supported: XLSX, XLS, CSV",type=["xlsx","xls","csv"],accept_multiple_files=True,key="excel_upload")
    if uploaded and st.button("Process Excel / CSV",type="primary"):
        all_frames=[]
        for f in uploaded:
            try:
                df=pd.read_csv(f) if f.name.lower().endswith(".csv") else extract_excel(f);result=process_tabular(df)
                if not result.empty:result["Source File"]=f.name;all_frames.append(result)
            except Exception as e:st.error(f"{f.name}: {e}")
        if all_frames:
            final=pd.concat(all_frames,ignore_index=True);grouped=final.groupby(["Customer","FG Code","Product"],dropna=False)["SAP Qty (PKT)"].sum().reset_index();st.success(f"Processed {len(grouped)} customer/product lines.");st.dataframe(grouped,use_container_width=True)
            st.subheader("Customer-wise SAP Orders");customers=grouped["Customer"].fillna("").replace("nan","").unique();txt_parts=[]
            for cust in customers:
                cust_df=grouped[grouped["Customer"].fillna("")==cust];title=cust if cust else "UNKNOWN CUSTOMER";st.markdown(f"#### {title}");block="\n".join(sap_line(row["FG Code"],row["SAP Qty (PKT)"]) for _,row in cust_df.iterrows());st.code(block,language="text");txt_parts.append(f"### {title}\n{block}")
            st.download_button("Download All Customer Orders","\n\n".join(txt_parts),file_name="SAP_Orders_All_Customers.txt",mime="text/plain")
        else:st.warning("No valid order rows found.")

st.divider();st.caption("SAP format: FG CODE + 2 tabs + QTY in PKT + 5 tabs + HO-WH + 2 tabs + CHEESE")