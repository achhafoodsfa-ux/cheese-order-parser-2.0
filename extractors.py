"""File readers used by the Streamlit UI (images, PDF, Word, Excel, CSV, text).

Kept separate from core.py so the parsing logic stays dependency free.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Tuple

IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "bmp", "tif", "tiff"}
TABULAR_EXTS = {"xlsx", "xls", "csv"}
TEXT_EXTS = {"txt"}
PDF_EXTS = {"pdf"}
DOC_EXTS = {"docx"}
ALL_EXTS = IMAGE_EXTS | TABULAR_EXTS | TEXT_EXTS | PDF_EXTS | DOC_EXTS


def clean_ocr_text(text: str) -> str:
    text = str(text or "").replace("\x00", " ")
    lines = []
    for raw in text.splitlines():
        line = re.sub(r"[ \t]+", " ", raw).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def tesseract_available() -> bool:
    try:
        import pytesseract

        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def ocr_image(data: bytes) -> str:
    """OCR a screenshot/photo with a light upscale + contrast pass."""
    from PIL import Image, ImageFilter, ImageOps
    import pytesseract

    image = Image.open(io.BytesIO(data))
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")
    width, height = image.size
    if max(width, height) < 2600:
        scale = 2600 / max(width, height)
        image = image.resize((int(width * scale), int(height * scale)), Image.Resampling.LANCZOS)
    gray = ImageOps.autocontrast(ImageOps.grayscale(image)).filter(ImageFilter.SHARPEN)
    text = pytesseract.image_to_string(gray, config="--psm 6")
    if len(re.sub(r"\s", "", text)) < 8:
        text = pytesseract.image_to_string(gray, config="--psm 11")
    return clean_ocr_text(text)


def pdf_to_text(data: bytes) -> str:
    from pypdf import PdfReader

    parts = []
    try:
        reader = PdfReader(io.BytesIO(data))
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
    except Exception:
        parts = []
    text = clean_ocr_text("\n".join(parts))
    if len(re.sub(r"\s", "", text)) >= 20:
        return text
    # Scanned PDF -> render each page and OCR it.
    try:
        import fitz  # PyMuPDF

        document = fitz.open(stream=data, filetype="pdf")
        pages = []
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            pages.append(ocr_image(pixmap.tobytes("png")))
        document.close()
        return clean_ocr_text("\n".join(pages))
    except Exception:
        return text


def docx_to_text(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    chunks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            chunks.append(" ".join(cell.text.strip() for cell in row.cells))
    return clean_ocr_text("\n".join(chunks))


def extract_excel(file):
    import pandas as pd

    workbook = pd.ExcelFile(file)
    frames = []
    for sheet in workbook.sheet_names:
        frame = workbook.parse(sheet)
        if not frame.empty:
            frame["__sheet__"] = sheet
            frames.append(frame)
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()


def read_csv(data: bytes):
    import pandas as pd

    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return pd.read_csv(io.BytesIO(data), encoding=encoding)
        except UnicodeDecodeError:
            continue
        except Exception:
            break
    return pd.read_csv(io.BytesIO(data), encoding="latin-1", engine="python", on_bad_lines="skip")


def read_uploaded_file(uploaded_file) -> Tuple[str, object]:
    """Return (kind, content) where kind is image/pdf/docx/csv/excel/text."""
    ext = Path(uploaded_file.name).suffix.lower().lstrip(".")
    data = uploaded_file.getvalue()
    if ext in IMAGE_EXTS:
        return "image", data
    if ext in PDF_EXTS:
        return "pdf", pdf_to_text(data)
    if ext in DOC_EXTS:
        return "docx", docx_to_text(data)
    if ext == "csv":
        return "csv", read_csv(data)
    if ext in ("xlsx", "xls"):
        return "excel", extract_excel(io.BytesIO(data))
    if ext in TEXT_EXTS:
        return "text", data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: .{ext}")
